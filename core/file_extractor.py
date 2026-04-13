"""
core/file_extractor.py — In-memory file text extraction for the public API.

Supports: PDF, images (OCR), plain text, CSV, JSON, Markdown, Word (.docx).
File bytes are read entirely in memory — nothing is written to disk.
"""
from __future__ import annotations

import io
import logging
import os

import pdfplumber
import pytesseract
from PIL import Image

logger = logging.getLogger(__name__)

_TEXT_EXTS  = {".txt", ".csv", ".json", ".md", ".log", ".yaml", ".yml", ".xml", ".html"}
_PDF_EXTS   = {".pdf"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".tif"}
_DOCX_EXTS  = {".docx", ".doc"}
_XLSX_EXTS  = {".xlsx"}
_XLS_EXTS   = {".xls"}

ALL_SUPPORTED_EXTS = _TEXT_EXTS | _PDF_EXTS | _IMAGE_EXTS | _DOCX_EXTS | _XLSX_EXTS | _XLS_EXTS


def extract_text_from_file(file_storage) -> str:
    """
    Accept a Flask FileStorage object and return its text content.

    Args:
        file_storage: werkzeug.datastructures.FileStorage

    Returns:
        Extracted text string (may be empty for blank files).

    Raises:
        ValueError: if the file extension is not supported.
    """
    filename = (file_storage.filename or "").strip()
    ext      = os.path.splitext(filename)[1].lower()
    mime     = (file_storage.mimetype or "").lower()

    data = file_storage.read()
    logger.debug("Extracting file: %s  ext=%s  mime=%s  size=%d bytes",
                 filename, ext, mime, len(data))

    # ── Plain text (by extension or MIME) ─────────────────────────────────────
    if ext in _TEXT_EXTS or mime.startswith("text/"):
        return data.decode("utf-8", errors="replace").strip()

    # ── PDF ───────────────────────────────────────────────────────────────────
    if ext in _PDF_EXTS or mime == "application/pdf":
        return _from_pdf(data)

    # ── Image (OCR) ───────────────────────────────────────────────────────────
    if ext in _IMAGE_EXTS or mime.startswith("image/"):
        return _from_image(data)

    # ── Word document ─────────────────────────────────────────────────────────
    if ext in _DOCX_EXTS or "wordprocessingml" in mime or "msword" in mime:
        return _from_docx(data, ext)

    # ── Excel spreadsheet ────────────────────────────────────────────────────
    if ext in _XLSX_EXTS or "spreadsheetml" in mime:
        return _from_xlsx(data)

    if ext in _XLS_EXTS or "ms-excel" in mime or "vnd.ms-excel" in mime:
        return _from_xls(data)

    raise ValueError(
        f"Unsupported file type: '{ext or mime}'. "
        f"Supported: PDF, images (PNG/JPG/TIFF/…), TXT, CSV, JSON, DOCX, DOC, XLSX, XLS."
    )


# ── Private helpers ────────────────────────────────────────────────────────────

def _from_pdf(data: bytes) -> str:
    pages = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
    return "\n\n".join(pages) if pages else "[PDF has no extractable text]"


def _from_image(data: bytes) -> str:
    img  = Image.open(io.BytesIO(data))
    text = pytesseract.image_to_string(img).strip()
    return text if text else "[No text found in image]"


def _from_docx(data: bytes, ext: str = ".docx") -> str:
    try:
        import docx
        doc   = docx.Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cell_texts:
                    lines.append(" | ".join(cell_texts))
        return "\n".join(lines) if lines else "[Word document has no extractable text]"
    except ImportError:
        return (
            "[python-docx is not installed — run: pip install python-docx "
            "to enable .docx extraction]"
        )
    except Exception as exc:
        return f"[Could not read Word document: {exc}]"


def _from_xlsx(data: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        rows_out = []
        for sheet in wb.worksheets:
            rows_out.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                line  = " | ".join(cells).strip(" |")
                if line:
                    rows_out.append(line)
        wb.close()
        return "\n".join(rows_out) if rows_out else "[Spreadsheet has no extractable text]"
    except ImportError:
        return "[openpyxl is not installed — run: pip install openpyxl to enable .xlsx extraction]"
    except Exception as exc:
        return f"[Could not read Excel file: {exc}]"


def _from_xls(data: bytes) -> str:
    try:
        import xlrd
        wb = xlrd.open_workbook(file_contents=data)
        rows_out = []
        for sheet in wb.sheets():
            rows_out.append(f"[Sheet: {sheet.name}]")
            for r in range(sheet.nrows):
                cells = [str(sheet.cell_value(r, c)) for c in range(sheet.ncols)]
                line  = " | ".join(cells).strip(" |")
                if line:
                    rows_out.append(line)
        return "\n".join(rows_out) if rows_out else "[Spreadsheet has no extractable text]"
    except ImportError:
        return "[xlrd is not installed — run: pip install xlrd to enable .xls extraction]"
    except Exception as exc:
        return f"[Could not read Excel file: {exc}]"
